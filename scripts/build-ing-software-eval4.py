#!/usr/bin/env python3
"""Genera docs/ing-software/eval-4-miguel-figuera.docx a partir de
docs/ing-software/eval-4-miguel-figuera.md

Formato: APA 7 (estudiante) en Arial, interlineado doble, margenes 2.54 cm,
numero de pagina arriba a la derecha, cuerpo JUSTIFICADO. CON portada e
INDICE con numeros de pagina reales (dos pasadas: docx -> PDF via LibreOffice
-> deteccion de paginas con pdftotext -> regeneracion). Soporta encabezados,
listas, **negrita**, tablas markdown y bloques de codigo ```...```.
El marcador [[INDICE]] en el .md indica donde va el indice.
"""
import re
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER,
)
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "ing-software" / "eval-4-miguel-figuera.md"
OUT = ROOT / "docs" / "ing-software" / "eval-4-miguel-figuera.docx"

FONT = "Arial"
MONO = "Consolas"
CONTENT_WIDTH_CM = 16.51  # carta 21.59 cm - margenes 2 x 2.54 cm

# Portada con el formato institucional del equipo (mismo de Taller 2 /
# Microtaller 3): membrete centrado, título al tercio superior, integrantes
# con C.I. a la derecha, tutor y ciudad-fecha centrada abajo. Sin numero de
# pagina en la portada.
PORTADA = {
    "membrete": [
        "REPÚBLICA BOLIVARIANA DE VENEZUELA",
        "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN SUPERIOR",
        "UNIVERSIDAD NACIONAL EXPERIMENTAL",
        "DE LAS TELECOMUNICACIONES E INFORMÁTICA",
    ],
    "titulo": "Evaluación 4: Informe Técnico — Procedimientos de Verificación del Producto de Software",
    "subtitulo": "Documento N.° PW-INF-04, versión 1.0",
    "integrantes": [
        "Miguel Figuera C.I: 23.558.789",
        "Iromy León C.I: V-30.243.131",
        "Alejandra Herde C.I: V-23.711.974",
        "Tutor: Misley Baute",
    ],
    "fecha": "Caracas, Julio de 2026",
}


def set_font(run, size=12, bold=False, mono=False, color=None):
    name = MONO if mono else FONT
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)


def add_runs_with_bold(par, text, size=12, base_bold=False):
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if chunk == "":
            continue
        is_bold = base_bold or (i % 2 == 1)
        set_font(par.add_run(chunk), size=size, bold=is_bold)


def configure_normal(doc):
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(12)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def set_margins(section):
    for side in ('top', 'bottom', 'left', 'right'):
        setattr(section, f"{side}_margin", Cm(2.54))


def add_page_number_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._element.append(f1); run._element.append(instr); run._element.append(f2)
    set_font(run, size=12)


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _blank_lines(doc, n):
    for _ in range(n):
        doc.add_paragraph()


def add_portada(doc):
    for line in PORTADA["membrete"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(line), size=12, bold=True)
    _blank_lines(doc, 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(PORTADA["titulo"]), size=14, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(PORTADA["subtitulo"]), size=10)
    _blank_lines(doc, 4)
    for line in PORTADA["integrantes"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(line), size=12, bold=True)
    _blank_lines(doc, 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(PORTADA["fecha"]), size=12, bold=True)
    add_page_break(doc)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), size=14 if level == 1 else 12, bold=True)


def add_body_paragraph(doc, text, justify=True, references=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if references:
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
    add_runs_with_bold(p, text)


def add_list_item(doc, text, ordered=False):
    style = 'List Number' if ordered else 'List Bullet'
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        text = ("• " + text) if not ordered else text
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs_with_bold(p, text)


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    for j, ln in enumerate(code_lines):
        if j > 0:
            run = p.add_run(); run.add_break()
        set_font(p.add_run(ln if ln else " "), size=10, mono=True, color=RGBColor(0x20, 0x20, 0x20))


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    for j, cell in enumerate(header):
        c = table.rows[0].cells[j]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_runs_with_bold(c.paragraphs[0], cell.strip(), size=10, base_bold=True)
    for r in body:
        cells = table.add_row().cells
        for j, cell in enumerate(r):
            if j < len(cells):
                cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[j].paragraphs[0].paragraph_format.line_spacing = 1.0
                add_runs_with_bold(cells[j].paragraphs[0], cell.strip(), size=10)


def split_table_row(line):
    return [c for c in line.strip().strip('|').split('|')]


def collect_headings(md):
    """Encabezados (nivel, texto) en orden de aparicion, para el indice."""
    out = []
    in_code = False
    for line in md.splitlines():
        s = line.strip()
        if s.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        if line.startswith("## "):
            out.append((2, line[3:].strip()))
        elif line.startswith("# "):
            out.append((1, line[2:].strip()))
    return out


def add_indice(doc, headings, page_map):
    add_heading(doc, "Índice", level=1)
    for level, text in headings:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        if level == 2:
            pf.left_indent = Cm(0.75)
        pf.tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        set_font(p.add_run(text), size=12, bold=(level == 1))
        p.add_run("\t")
        page = page_map.get(text, "—") if page_map else "—"
        set_font(p.add_run(str(page)), size=12)
    add_page_break(doc)


def parse_markdown(doc, md, headings, page_map):
    lines = md.splitlines()
    in_refs = False
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped == "":
            i += 1
            continue
        if stripped == "[[INDICE]]":
            add_page_break(doc)
            add_indice(doc, headings, page_map)
            i += 1
            continue
        if stripped.startswith("```"):
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, code)
            continue
        if stripped.startswith("|"):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not re.match(r"^\|[\s:\-|]+\|?$", row):
                    tbl.append(split_table_row(row))
                i += 1
            if tbl:
                add_table(doc, tbl)
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("# "):
            title = line[2:].strip()
            in_refs = title.lower().startswith("referencias")
            add_heading(doc, title, level=1)
        elif line.startswith("- "):
            add_list_item(doc, line[2:].strip(), ordered=False)
        elif re.match(r"^\d+\.\s", line):
            add_list_item(doc, re.sub(r"^\d+\.\s", "", line).strip(), ordered=True)
        else:
            add_body_paragraph(doc, stripped, justify=not in_refs, references=in_refs)
        i += 1


def build(md, headings, page_map, out_path):
    doc = Document()
    configure_normal(doc)
    sec = doc.sections[0]
    set_margins(sec)
    sec.different_first_page_header_footer = True  # portada sin numero de pagina
    add_page_number_header(sec)
    add_portada(doc)
    parse_markdown(doc, md, headings, page_map)
    doc.save(str(out_path))


def detect_pages(docx_path, headings):
    """Convierte a PDF y ubica la pagina de cada encabezado (ignorando la
    pagina del indice, donde los titulos tambien aparecen como entradas)."""
    with tempfile.TemporaryDirectory() as tmp:
        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", tmp, str(docx_path)],
            check=True, capture_output=True,
        )
        pdf = Path(tmp) / (Path(docx_path).stem + ".pdf")
        txt = subprocess.run(
            ["pdftotext", "-layout", str(pdf), "-"],
            check=True, capture_output=True, text=True,
        ).stdout
    pages = txt.split("\f")

    def page_lines(pg):
        return [ln.strip() for ln in pg.splitlines() if ln.strip()]

    indice_page = next(
        (idx for idx, pg in enumerate(pages, start=1) if "Índice" in page_lines(pg)),
        None,
    )
    page_map = {}
    pos = 1
    for _, text in headings:
        for idx in range(pos, len(pages) + 1):
            if idx == indice_page:
                continue
            if text in page_lines(pages[idx - 1]):
                page_map[text] = idx
                pos = idx
                break
    return page_map


def main():
    md = SRC.read_text(encoding="utf-8")
    headings = collect_headings(md)
    # Pasada 1: sin numeros de pagina, solo para medir la paginacion real.
    build(md, headings, None, OUT)
    page_map = detect_pages(OUT, headings)
    missing = [t for _, t in headings if t not in page_map]
    if missing:
        print(f"AVISO: encabezados sin pagina detectada: {missing}")
    # Pasada 2: indice definitivo con las paginas medidas.
    build(md, headings, page_map, OUT)
    print(f"OK -> {OUT}")
    for lvl, t in headings:
        print(f"  {'  ' if lvl == 2 else ''}{t} ... pag. {page_map.get(t, '?')}")


if __name__ == "__main__":
    main()
