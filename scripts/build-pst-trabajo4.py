#!/usr/bin/env python3
"""Genera docs/pst/trabajo4-informe-final.docx a partir de docs/pst/trabajo4-informe-final.md
Formato: APA 7 (estudiante) en Arial, interlineado doble, margenes 2.54 cm,
numero de pagina arriba a la derecha, con el cuerpo JUSTIFICADO.
CON portada institucional (formato del equipo: membrete, titulo, integrantes
con C.I., tutora y ciudad-fecha; sin numero de pagina en la portada).
Soporta encabezados, listas, **negrita**, tablas markdown y bloques ```...```.

Anexos (interlineado compacto):
  A. Requisitos funcionales y no funcionales (docs/mobile-requirements.md)
  B. Demo de interfaz (docs/PawCare Mobile — Demo UI.pdf → imagenes, requiere pdftoppm)
  C. Diagramas de arquitectura (docs/PawCare Mobile — Diagramas de Arquitectura.pdf → imagenes)
  D. Rutas del API Rails (docs/rails-routes.txt como tabla)
"""
import re
import subprocess
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "pst" / "presentacion" / "trabajo4-informe-final.md"
OUT = ROOT / "docs" / "pst" / "presentacion" / "trabajo4-informe-final.docx"

DOCS = ROOT / "docs"
ANEXO_REQ_MD = DOCS / "mobile-requirements.md"
ANEXO_DEMO_PDF = DOCS / "PawCare Mobile — Demo UI.pdf"
ANEXO_DIAG_PDF = DOCS / "PawCare Mobile — Diagramas de Arquitectura.pdf"
ANEXO_ROUTES = DOCS / "rails-routes.txt"

# ancho util de pagina A4 con margenes de 2.54 cm
CONTENT_WIDTH_CM = 15.92

FONT = "Arial"
MONO = "Consolas"

PORTADA = {
    "membrete": [
        "REPÚBLICA BOLIVARIANA DE VENEZUELA",
        "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN SUPERIOR",
        "UNIVERSIDAD NACIONAL EXPERIMENTAL",
        "DE LAS TELECOMUNICACIONES E INFORMÁTICA",
    ],
    "titulo": "Informe Final: Pawcare — App Móvil para Atención Veterinaria a Domicilio",
    "integrantes": [
        "Miguel Figuera C.I: 23.558.789",
        "Iromy León C.I: V-30.243.131",
        "Alejandra Herde C.I: V-23.711.974",
        "Tutora: Yuly Delgado",
    ],
    "fecha": "La Victoria, Julio de 2026",
}


def set_font(run, size=12, bold=False, mono=False, color=None):
    name = MONO if mono else FONT
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), name)


def add_runs_with_bold(par, text, size=12, base_bold=False):
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if chunk == "":
            continue
        is_bold = base_bold or (i % 2 == 1)
        set_font(par.add_run(chunk), size=size, bold=is_bold)


def configure_normal(doc):
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(12)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def set_margins(section):
    for side in ('top', 'bottom', 'left', 'right'):
        setattr(section, f"{side}_margin", Cm(2.54))


def add_page_number_header(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._element.append(f1); run._element.append(instr); run._element.append(f2)
    set_font(run, size=12)


def _blank_lines(doc, n):
    for _ in range(n):
        doc.add_paragraph()


def add_portada(doc):
    for line in PORTADA["membrete"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(line), size=12, bold=True)
    _blank_lines(doc, 5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(PORTADA["titulo"]), size=14, bold=True)
    _blank_lines(doc, 4)
    for line in PORTADA["integrantes"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(line), size=12, bold=True)
    _blank_lines(doc, 4)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(PORTADA["fecha"]), size=12, bold=True)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc, text, level, compact=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if compact:
        p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), size=14 if level == 1 else 12, bold=True)


def add_body_paragraph(doc, text, justify=True, references=False, compact=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if references:
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
    if compact:
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(4)
    add_runs_with_bold(p, text)


def add_list_item(doc, text, ordered=False, compact=False):
    style = 'List Number' if ordered else 'List Bullet'
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        text = ("• " + text) if not ordered else text
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if compact:
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(2)
    add_runs_with_bold(p, text)


def add_code_block(doc, code_lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # sombreado de fondo gris claro
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    for j, ln in enumerate(code_lines):
        if j > 0:
            run = p.add_run(); run.add_break()
        set_font(p.add_run(ln if ln else " "), size=10, mono=True, color=RGBColor(0x20, 0x20, 0x20))


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    for j, cell in enumerate(header):
        c = table.rows[0].cells[j]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_runs_with_bold(c.paragraphs[0], cell.strip(), size=10, base_bold=True)
    for r in body:
        cells = table.add_row().cells
        for j, cell in enumerate(r):
            if j < len(cells):
                cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                cells[j].paragraphs[0].paragraph_format.line_spacing = 1.0
                add_runs_with_bold(cells[j].paragraphs[0], cell.strip(), size=10)


def split_table_row(line):
    return [c for c in line.strip().strip('|').split('|')]


def parse_markdown(doc, md, compact=False):
    lines = md.splitlines()
    in_refs = False
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i].rstrip()
        stripped = line.strip()
        if stripped == "" or re.match(r"^-{3,}$", stripped):
            i += 1
            continue
        if stripped.startswith("> "):
            stripped = stripped[2:].strip()
            line = stripped
        # bloque de codigo
        if stripped.startswith("```"):
            code = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1  # cerrar ```
            add_code_block(doc, code)
            continue
        # tabla markdown
        if stripped.startswith("|"):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                row = lines[i].strip()
                if not re.match(r"^\|[\s:\-|]+\|?$", row):  # saltar separador ---
                    tbl.append(split_table_row(row))
                i += 1
            if tbl:
                add_table(doc, tbl)
            continue
        if line.startswith("#### "):
            add_body_paragraph(doc, "**" + line[5:].strip() + "**", justify=False, compact=compact)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=2, compact=compact)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2, compact=compact)
        elif line.startswith("# "):
            title = line[2:].strip()
            in_refs = title.lower().startswith("referencias")
            add_heading(doc, title, level=1, compact=compact)
        elif line.startswith("- "):
            add_list_item(doc, line[2:].strip(), ordered=False, compact=compact)
        elif re.match(r"^\d+\.\s", line):
            add_list_item(doc, re.sub(r"^\d+\.\s", "", line).strip(), ordered=True, compact=compact)
        else:
            add_body_paragraph(doc, stripped, justify=not in_refs, references=in_refs, compact=compact)
        i += 1


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_annex_heading(doc, letra, titulo, fuente=None):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(f"ANEXO {letra}"), size=14, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(titulo), size=12, bold=True)
    if fuente:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(f"Fuente: {fuente}"), size=10)


def add_annex_cover(doc):
    add_page_break(doc)
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("ANEXOS"), size=16, bold=True)
    doc.add_paragraph()
    for linea in (
        "Anexo A. Requisitos funcionales y no funcionales del cliente móvil",
        "Anexo B. Demo de interfaz de usuario (Demo UI)",
        "Anexo C. Diagramas de arquitectura",
        "Anexo D. Rutas del API Rails consumidas por la aplicación (rails-routes.txt)",
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        set_font(p.add_run(linea), size=12)


def add_pdf_pages_as_images(doc, pdf_path, dpi=130):
    """Rasteriza cada pagina del PDF con pdftoppm y la inserta a ancho completo."""
    with tempfile.TemporaryDirectory() as tmp:
        prefix = Path(tmp) / "page"
        subprocess.run(
            ["pdftoppm", "-png", "-r", str(dpi), str(pdf_path), str(prefix)],
            check=True,
        )
        pages = sorted(Path(tmp).glob("page-*.png"))
        if not pages:
            raise RuntimeError(f"pdftoppm no produjo paginas para {pdf_path}")
        for png in pages:
            doc.add_picture(str(png), width=Cm(CONTENT_WIDTH_CM))
            pic = doc.paragraphs[-1]
            pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic.paragraph_format.line_spacing = 1.0
            pic.paragraph_format.space_after = Pt(6)


VERBOS_HTTP = {"GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS", "HEAD"}


def parse_routes_line(line):
    tokens = line.split()
    if len(tokens) >= 4 and tokens[1] in VERBOS_HTTP:
        return tokens[0], tokens[1], tokens[2], " ".join(tokens[3:])
    if len(tokens) == 3:
        if tokens[0] in VERBOS_HTTP:
            return "", tokens[0], tokens[1], tokens[2]
        return tokens[0], "", tokens[1], tokens[2]
    if len(tokens) == 2:
        return "", "", tokens[0], tokens[1]
    return "", "", line.strip(), ""


def add_routes_table(doc, txt_path):
    lines = [ln for ln in txt_path.read_text(encoding="utf-8").splitlines() if ln.strip()]
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    anchos = (Cm(4.1), Cm(1.5), Cm(5.7), Cm(4.6))
    encabezado = ("Prefix", "Verb", "URI Pattern", "Controller#Action")
    for j, texto in enumerate(encabezado):
        c = table.rows[0].cells[j]
        c.width = anchos[j]
        c.paragraphs[0].paragraph_format.line_spacing = 1.0
        set_font(c.paragraphs[0].add_run(texto), size=8, bold=True)
    for ln in lines[1:]:  # salta la fila de encabezado del propio archivo
        celdas = table.add_row().cells
        for j, texto in enumerate(parse_routes_line(ln)):
            celdas[j].width = anchos[j]
            par = celdas[j].paragraphs[0]
            par.paragraph_format.line_spacing = 1.0
            set_font(par.add_run(texto), size=7, mono=True)


def add_annexes(doc):
    add_annex_cover(doc)

    add_annex_heading(
        doc, "A", "Requisitos funcionales y no funcionales del cliente móvil",
        fuente="docs/mobile-requirements.md — 26 RF y 17 RNF trazados a las 157 rutas del API",
    )
    parse_markdown(doc, ANEXO_REQ_MD.read_text(encoding="utf-8"), compact=True)

    add_annex_heading(
        doc, "B", "Demo de interfaz de usuario (Demo UI)",
        fuente="docs/PawCare Mobile — Demo UI.pdf — sistema visual con temas claro y oscuro",
    )
    add_pdf_pages_as_images(doc, ANEXO_DEMO_PDF)

    add_annex_heading(
        doc, "C", "Diagramas de arquitectura",
        fuente="docs/PawCare Mobile — Diagramas de Arquitectura.pdf (v4.0, julio 2026)",
    )
    add_pdf_pages_as_images(doc, ANEXO_DIAG_PDF)

    add_annex_heading(
        doc, "D", "Rutas del API Rails consumidas por la aplicación",
        fuente="docs/rails-routes.txt — generado con `rails routes` desde el backend",
    )
    add_routes_table(doc, ANEXO_ROUTES)


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    configure_normal(doc)
    sec = doc.sections[0]
    set_margins(sec)
    sec.different_first_page_header_footer = True  # portada sin numero de pagina
    add_page_number_header(sec)
    add_portada(doc)
    parse_markdown(doc, md)
    add_annexes(doc)
    doc.save(str(OUT))
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
