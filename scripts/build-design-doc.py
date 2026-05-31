#!/usr/bin/env python3
"""Genera docs/diseno-proyecto-pawcare.docx a partir de docs/diseno-proyecto-pawcare.md
Formato: APA 7 (estudiante) en Arial, interlineado doble, margenes 2.54cm,
con la excepcion de que el cuerpo va JUSTIFICADO (no flush-left).
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "docs" / "diseno-proyecto-pawcare.md"
OUT = ROOT / "docs" / "diseno-proyecto-pawcare.docx"

FONT = "Arial"

# ---------- datos de portada ----------
MEMBRETE = [
    "REPÚBLICA BOLIVARIANA DE VENEZUELA",
    "MINISTERIO DEL PODER POPULAR PARA LA EDUCACIÓN SUPERIOR",
    "UNIVERSIDAD NACIONAL EXPERIMENTAL",
    "DE LAS TELECOMUNICACIONES E INFORMÁTICA",
]
TITULO = "Pawcare: App móvil para atención veterinaria a domicilio"
INTEGRANTES = [
    "Miguel Figuera C.I: 23.558.789",
    "Iromy León C.I: V-30.243.131",
    "Alejandra Herde C.I: V-23.711.974",
]
TUTOR = "Tutor: Yuly Delgado"
FECHA = "Caracas, Mayo de 2026"


def set_font(run, size=12, bold=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    # asegurar Arial tambien para scripts complejos/east-asian
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), FONT)


def add_runs_with_bold(par, text, size=12, base_bold=False):
    """Divide el texto por **negrita** y agrega runs."""
    for i, chunk in enumerate(re.split(r"\*\*(.+?)\*\*", text)):
        if chunk == "":
            continue
        is_bold = base_bold or (i % 2 == 1)
        set_font(par.add_run(chunk), size=size, bold=is_bold)


def configure_normal(doc):
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(12)
    rpr = st.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
        rfonts.set(qn(attr), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)


def set_margins(section):
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def add_page_number_header(section):
    """Numero de pagina arriba a la derecha (campo PAGE)."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1); run._element.append(instr); run._element.append(fldChar2)
    set_font(run, size=12)


def blank(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0


def build_cover(doc):
    # membrete (centrado, negrita)
    for line in MEMBRETE:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        set_font(p.add_run(line), size=12, bold=True)
    blank(doc, 7)
    # titulo (centrado, negrita, grande)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(TITULO), size=16, bold=True)
    blank(doc, 8)
    # integrantes (derecha, negrita)
    for line in INTEGRANTES:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.line_spacing = 1.15
        set_font(p.add_run(line), size=12, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(TUTOR), size=12, bold=True)
    blank(doc, 6)
    # lugar y fecha (centrado)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(FECHA), size=12, bold=True)
    # salto de pagina
    doc.add_page_break()


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), size=14 if level == 1 else 12, bold=True)
    return p


def add_body_paragraph(doc, text, justify=True, references=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if references:
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
    add_runs_with_bold(p, text)
    return p


def add_list_item(doc, text, ordered=False):
    style = 'List Number' if ordered else 'List Bullet'
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
        text = ("• " + text) if not ordered else text
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_runs_with_bold(p, text)
    return p


def parse_markdown(doc, md):
    lines = md.splitlines()
    in_refs = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.strip() == "":
            i += 1
            continue
        if line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("# "):
            title = line[2:].strip()
            in_refs = title.lower().startswith("referencias")
            add_heading(doc, title, level=1)
        elif line.startswith("- "):
            add_list_item(doc, line[2:].strip(), ordered=False)
        elif re.match(r"^\d+\.\s", line):
            add_list_item(doc, re.sub(r"^\d+\.\s", "", line).strip(), ordered=True)
        else:
            add_body_paragraph(doc, line.strip(), justify=not in_refs, references=in_refs)
        i += 1


def main():
    md = SRC.read_text(encoding="utf-8")
    doc = Document()
    configure_normal(doc)
    sec = doc.sections[0]
    set_margins(sec)
    add_page_number_header(sec)
    build_cover(doc)
    parse_markdown(doc, md)
    doc.save(str(OUT))
    print(f"OK -> {OUT}")


if __name__ == "__main__":
    main()
